# -*- coding: utf-8 -*-
"""
知识库解析与检索（多模态 P1 版）
- 支持类型：FAQ / 文档 (pdf/docx/txt/md) / URL / 文本笔记
- 切块：~500 字一段，重叠 ~80 字
- 检索：BM25（基于 jieba 分词），Top-K 召回
"""
import json
import math
import os
import re
from typing import List, Dict, Tuple, Optional

import requests

# ============================================================
#  分词（jieba 优先；缺失时回退正则切分）
# ============================================================
try:
    import jieba
    jieba.initialize()
    _HAS_JIEBA = True
except Exception:
    _HAS_JIEBA = False


_STOPWORDS = set("""的 了 是 我 你 他 她 它 们 在 有 和 与 及 或 但 而 也 都 就 还 让 给 把 被 等 这 那
之 吗 啊 呢 哦 哈 嗯 一 个 是不是 怎么 怎样 如何 什么 哪个 哪些 多少 多 少 啊
the a an of is are be to in on for with at by from as it that this those these and or but
有些 一些 我们 你们 他们 一下 一直 一定 啊 噢""".split())


def tokenize(text: str) -> List[str]:
    """中文 + 英文分词，去停用词；返回小写 token 列表"""
    if not text:
        return []
    text = text.lower()
    if _HAS_JIEBA:
        words = [w.strip() for w in jieba.lcut(text) if w.strip()]
    else:
        # 回退：中文按字 + 英文按词
        words = re.findall(r"[a-z0-9]+|[\u4e00-\u9fa5]", text)
    return [w for w in words if w and w not in _STOPWORDS and len(w) >= 1]


# ============================================================
#  切块
# ============================================================
def split_chunks(text: str, max_chars: int = 500, overlap: int = 80) -> List[str]:
    """按 ~max_chars 切块，相邻块有少量重叠以保留上下文"""
    if not text:
        return []
    text = re.sub(r"\r\n", "\n", text).strip()
    # 优先按段落切
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", text) if p.strip()]
    chunks: List[str] = []
    buf = ""
    for p in paragraphs:
        if len(buf) + len(p) + 1 <= max_chars:
            buf = (buf + "\n" + p) if buf else p
        else:
            if buf:
                chunks.append(buf)
            # 单段过长 → 按句号拆
            if len(p) > max_chars:
                sentences = re.split(r"(?<=[。！？!?；;])\s*", p)
                cur = ""
                for s in sentences:
                    if not s:
                        continue
                    if len(cur) + len(s) <= max_chars:
                        cur += s
                    else:
                        if cur:
                            chunks.append(cur)
                        if len(s) > max_chars:
                            # 强制切
                            for i in range(0, len(s), max_chars):
                                chunks.append(s[i:i + max_chars])
                            cur = ""
                        else:
                            cur = s
                if cur:
                    buf = cur
                else:
                    buf = ""
            else:
                buf = p
    if buf:
        chunks.append(buf)

    # 加重叠（让相邻块尾部出现在下一块开头）
    if overlap > 0 and len(chunks) > 1:
        out = [chunks[0]]
        for i in range(1, len(chunks)):
            prev_tail = chunks[i - 1][-overlap:]
            out.append(prev_tail + "\n" + chunks[i])
        chunks = out
    return chunks


# ============================================================
#  解析器：返回 (full_text, page_pieces) 形式
#  page_pieces: List[(page_no, text_of_page)]; 没分页就 [(None, full_text)]
# ============================================================
def parse_pdf(file_path: str) -> List[Tuple[Optional[int], str]]:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # 兼容
        except ImportError:
            raise RuntimeError("需要安装 pypdf：pip install pypdf")
    reader = PdfReader(file_path)
    pieces: List[Tuple[Optional[int], str]] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        txt = txt.strip()
        if txt:
            pieces.append((i, txt))
    return pieces


def parse_docx(file_path: str) -> List[Tuple[Optional[int], str]]:
    try:
        import docx  # python-docx
    except ImportError:
        raise RuntimeError("需要安装 python-docx：pip install python-docx")
    doc = docx.Document(file_path)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    # 也读表格
    for table in doc.tables:
        for row in table.rows:
            row_txt = " | ".join((c.text or "").strip() for c in row.cells)
            if row_txt.strip():
                parts.append(row_txt)
    return [(None, "\n\n".join(parts))]


def parse_txt(file_path: str) -> List[Tuple[Optional[int], str]]:
    encs = ["utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"]
    for enc in encs:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return [(None, f.read())]
        except UnicodeDecodeError:
            continue
    with open(file_path, "rb") as f:
        return [(None, f.read().decode("utf-8", errors="ignore"))]


def parse_url(url: str, timeout: int = 15) -> Tuple[str, str]:
    """抓取网页 → (title, plain_text)"""
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0 Safari/537.36"
    }
    r = requests.get(url, headers=headers, timeout=timeout)
    r.raise_for_status()
    enc = r.encoding or "utf-8"
    if enc.lower() == "iso-8859-1":
        enc = r.apparent_encoding or "utf-8"
    html = r.content.decode(enc, errors="ignore")

    # 优先 BeautifulSoup
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        # 移除脚本/样式
        for tag in soup(["script", "style", "noscript", "iframe", "header", "footer", "nav"]):
            tag.decompose()
        title = (soup.title.string.strip() if soup.title and soup.title.string else url)
        # 优先取 <article> / <main>，再回退到 body
        main = soup.find("article") or soup.find("main") or soup.body or soup
        text = main.get_text("\n", strip=True)
    except Exception:
        title_match = re.search(r"<title[^>]*>(.*?)</title>", html, re.S | re.I)
        title = title_match.group(1).strip() if title_match else url
        text = re.sub(r"<script.*?</script>|<style.*?</style>", "", html, flags=re.S | re.I)
        text = re.sub(r"<[^>]+>", "\n", text)
        text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text).strip()
    return title, text


def parse_doc_to_chunks(doc_type: str, file_path: str = None,
                       url: str = None, raw_text: str = None,
                       title_hint: str = None) -> Tuple[str, List[Dict]]:
    """统一入口：根据类型解析 → 返回 (resolved_title, [chunk_dict])
    chunk_dict: {content, page_no, section, chunk_index}
    """
    title = title_hint or ""
    chunks: List[Dict] = []

    if doc_type in ("note", "faq"):
        text = raw_text or ""
        if not title:
            title = (text.split("\n")[0] or "笔记")[:80]
        for i, c in enumerate(split_chunks(text)):
            chunks.append({"content": c, "page_no": None, "section": None, "chunk_index": i})

    elif doc_type == "url":
        if not url:
            raise ValueError("URL 类型需要提供 url")
        new_title, text = parse_url(url)
        if not title:
            title = new_title
        for i, c in enumerate(split_chunks(text)):
            chunks.append({"content": c, "page_no": None, "section": None, "chunk_index": i})

    elif doc_type == "doc":
        if not file_path or not os.path.exists(file_path):
            raise ValueError(f"文件不存在: {file_path}")
        ext = file_path.lower().rsplit(".", 1)[-1] if "." in file_path else ""
        if ext == "pdf":
            pieces = parse_pdf(file_path)
        elif ext in ("docx",):
            pieces = parse_docx(file_path)
        elif ext in ("txt", "md", "markdown"):
            pieces = parse_txt(file_path)
        else:
            raise ValueError(f"暂不支持的文件类型 .{ext}（P2 支持图片/音频）")
        idx = 0
        for page_no, piece_text in pieces:
            for c in split_chunks(piece_text):
                chunks.append({
                    "content": c,
                    "page_no": page_no,
                    "section": None,
                    "chunk_index": idx,
                })
                idx += 1
        if not title:
            title = os.path.splitext(os.path.basename(file_path))[0][:80]

    elif doc_type == "image":
        if not file_path or not os.path.exists(file_path):
            raise ValueError(f"图片文件不存在: {file_path}")
        # 调视觉模型生成详细描述
        desc = describe_image(file_path)
        if not desc:
            raise RuntimeError("视觉模型未返回描述（请检查 GLM-4V-Flash 配置）")
        if not title:
            title = os.path.splitext(os.path.basename(file_path))[0][:80]
        # 一张图通常一个 chunk 即可
        for i, c in enumerate(split_chunks(desc, max_chars=600, overlap=0)):
            chunks.append({
                "content": c, "page_no": None, "section": "图片描述",
                "chunk_index": i,
            })

    else:
        raise ValueError(f"暂不支持的类型: {doc_type}（audio 在 P3 阶段开放）")

    return title, chunks


# ============================================================
#  图片视觉理解（GLM-4V-Flash 免费）
# ============================================================
def _read_image_as_b64(file_path: str, max_side: int = 1280) -> Tuple[str, str]:
    """读取图片为 base64；过大则缩放到 max_side 以内（节省传输）
    返回 (mime, base64)
    """
    try:
        from PIL import Image
        from io import BytesIO
        import base64
        img = Image.open(file_path)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        w, h = img.size
        if max(w, h) > max_side:
            ratio = max_side / max(w, h)
            img = img.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
        buf = BytesIO()
        img.save(buf, "JPEG", quality=85)
        b = buf.getvalue()
        return "image/jpeg", base64.b64encode(b).decode()
    except ImportError:
        # 没有 Pillow 就直接读原文件
        import base64
        ext = file_path.lower().rsplit(".", 1)[-1] if "." in file_path else "jpeg"
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "webp": "image/webp", "gif": "image/gif"}.get(ext, "image/jpeg")
        with open(file_path, "rb") as f:
            return mime, base64.b64encode(f.read()).decode()


def _get_vision_provider():
    """获取一个可用的视觉模型 Provider：
    优先级：用户在 AIProvider 配置了 glm-4v 系列 → 复用其 key；
    否则借用同 base_url 的 GLM-4-Flash 的 key（智谱平台同 key 通用）
    """
    from models import AIProvider
    # 1) 显式视觉模型
    p = AIProvider.query.filter(
        AIProvider.enabled.is_(True),
        AIProvider.model.like("%glm-4v%"),
    ).first()
    if p:
        return p.api_key, p.base_url, p.model
    # 2) 借用智谱普通模型的 Key
    p = AIProvider.query.filter(
        AIProvider.enabled.is_(True),
        AIProvider.base_url.like("%bigmodel.cn%"),
    ).first()
    if p:
        return p.api_key, p.base_url, "glm-4v-flash"
    return None, None, None


VISION_PROMPT = """请详细描述这张图片的内容，重点关注：
1. 如果是文档/截图/聊天记录：提取所有可见的文字（OCR），保持原格式排版
2. 如果是证件/材料：识别证件类型、关键字段（姓名、编号、有效期等，模糊处理身份证号最后 4 位）
3. 如果是地图/标志：识别地点名称、机构名称
4. 如果是普通照片：描述场景、物品、人物动作
5. 涉及青年驿站、人才补贴、求职材料的内容尤其要详细记录

输出格式：
[图片类型] xxx
[关键文字] xxx
[详细描述] xxx
（用中文回答，越详细越好，便于后续检索）"""


def describe_image(file_path: str, custom_prompt: str = None) -> str:
    """调用视觉模型描述图片，返回纯文本描述"""
    api_key, base_url, model = _get_vision_provider()
    if not api_key:
        raise RuntimeError("未找到可用的视觉模型 Provider（请先配置智谱 GLM Key）")
    mime, b64 = _read_image_as_b64(file_path)

    payload = {
        "model": model or "glm-4v-flash",
        "messages": [{
            "role": "user",
            "content": [
                {"type": "text", "text": custom_prompt or VISION_PROMPT},
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
            ],
        }],
        "temperature": 0.2,
        "max_tokens": 800,
    }
    r = requests.post(
        (base_url or "https://open.bigmodel.cn/api/paas/v4").rstrip("/") + "/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json=payload, timeout=60,
    )
    if r.status_code != 200:
        raise RuntimeError(f"视觉模型 HTTP {r.status_code}: {r.text[:300]}")
    data = r.json()
    return (data.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()


# ============================================================
#  BM25 检索
# ============================================================
class BM25Index:
    def __init__(self, k1: float = 1.5, b: float = 0.75):
        self.k1 = k1
        self.b = b
        self.docs: List[List[str]] = []     # [tokens]
        self.meta: List[dict] = []           # [extra info per chunk]
        self.df: Dict[str, int] = {}
        self.idf: Dict[str, float] = {}
        self.doc_lens: List[int] = []
        self.avgdl: float = 0.0

    def add(self, tokens: List[str], meta: dict):
        if not tokens:
            return
        self.docs.append(tokens)
        self.meta.append(meta)
        self.doc_lens.append(len(tokens))
        for t in set(tokens):
            self.df[t] = self.df.get(t, 0) + 1

    def build(self):
        N = len(self.docs)
        if not N:
            return
        self.avgdl = sum(self.doc_lens) / N
        for term, df in self.df.items():
            # BM25 IDF（避免负值）
            self.idf[term] = math.log(1 + (N - df + 0.5) / (df + 0.5))

    def search(self, query: str, top_k: int = 5) -> List[Tuple[float, dict]]:
        if not self.docs:
            return []
        q_tokens = tokenize(query)
        if not q_tokens:
            return []
        scores: List[float] = [0.0] * len(self.docs)
        # 词频（每个 doc）
        for i, doc_tokens in enumerate(self.docs):
            if not doc_tokens:
                continue
            dl = self.doc_lens[i]
            tf: Dict[str, int] = {}
            for w in doc_tokens:
                tf[w] = tf.get(w, 0) + 1
            score = 0.0
            for q in q_tokens:
                if q not in tf:
                    continue
                f = tf[q]
                idf = self.idf.get(q, 0.0)
                score += idf * (f * (self.k1 + 1)) / (f + self.k1 * (1 - self.b + self.b * dl / max(self.avgdl, 1)))
            scores[i] = score
        ranked = sorted(zip(scores, self.meta), key=lambda x: x[0], reverse=True)
        return [(s, m) for s, m in ranked[:top_k] if s > 0]


def search_chunks(query: str, top_k: int = 5, max_per_doc: int = 2) -> List[Dict]:
    """对所有启用的 KnowledgeChunk 做 BM25 检索，返回:
    [{score, doc_id, chunk_id, content, page_no, doc_title, doc_type, summary}]
    每个文档最多保留 max_per_doc 个 chunk（避免单文档霸屏）"""
    from models import KnowledgeChunk, KnowledgeDoc, db
    from sqlalchemy.orm import joinedload

    rows = (KnowledgeChunk.query
            .options(joinedload(KnowledgeChunk.doc))
            .filter(KnowledgeChunk.enabled.is_(True))
            .join(KnowledgeDoc, KnowledgeDoc.id == KnowledgeChunk.doc_id)
            .filter(KnowledgeDoc.enabled.is_(True))
            .all())
    if not rows:
        return []

    idx = BM25Index()
    for c in rows:
        # 优先使用预存的 tokens；否则即时分词
        toks = c.tokens.split() if c.tokens else tokenize(c.content)
        idx.add(toks, {
            "doc_id": c.doc_id, "chunk_id": c.id,
            "content": c.content, "page_no": c.page_no,
            "doc_title": c.doc.title if c.doc else "未命名",
            "doc_type": c.doc.doc_type if c.doc else "note",
            "summary": (c.doc.summary or "") if c.doc else "",
            "source_url": (c.doc.source_url or "") if c.doc else "",
        })
    idx.build()
    raw = idx.search(query, top_k=top_k * 3)

    # 按文档去重，保留 top_k 个最终结果
    final: List[Dict] = []
    per_doc_count: Dict[int, int] = {}
    for score, meta in raw:
        did = meta["doc_id"]
        if per_doc_count.get(did, 0) >= max_per_doc:
            continue
        per_doc_count[did] = per_doc_count.get(did, 0) + 1
        final.append({"score": round(score, 3), **meta})
        if len(final) >= top_k:
            break
    return final


# ============================================================
#  入库主流程：解析 → 切块 → 写库（带 token 预存）
# ============================================================
def ingest(doc_id: int):
    """异步/同步入库流程：根据 KnowledgeDoc.id 完成解析 + 切块 + token 预存"""
    from models import KnowledgeDoc, KnowledgeChunk, db
    doc = KnowledgeDoc.query.get(doc_id)
    if not doc:
        return False, "文档不存在"
    try:
        doc.status = "parsing"
        db.session.commit()

        kwargs = {
            "doc_type": doc.doc_type,
            "title_hint": doc.title,
        }
        if doc.doc_type == "url":
            kwargs["url"] = doc.source_url
        elif doc.doc_type == "doc":
            kwargs["file_path"] = doc.file_path
        elif doc.doc_type in ("note", "faq"):
            kwargs["raw_text"] = doc.summary or ""

        title, chunks = parse_doc_to_chunks(**kwargs)
        if title and (not doc.title or doc.title == "(待解析)"):
            doc.title = title[:300]

        # 清掉旧 chunk
        KnowledgeChunk.query.filter_by(doc_id=doc.id).delete(synchronize_session=False)

        for c in chunks:
            toks = tokenize(c["content"])
            kc = KnowledgeChunk(
                doc_id=doc.id,
                chunk_index=c.get("chunk_index", 0),
                content=c["content"],
                tokens=" ".join(toks),
                page_no=c.get("page_no"),
                section=c.get("section"),
                enabled=True,
            )
            db.session.add(kc)

        # 自动摘要：取前 200 字（或 FAQ 直接答案）
        if not doc.summary or doc.doc_type in ("doc", "url"):
            joined = "\n".join([c["content"] for c in chunks[:3]])
            doc.summary = (joined[:300]).strip()

        doc.chunk_count = len(chunks)
        doc.status = "ready"
        doc.error_msg = None
        db.session.commit()
        return True, f"解析完成，共生成 {len(chunks)} 个切块"
    except Exception as e:
        db.session.rollback()
        try:
            doc.status = "failed"
            doc.error_msg = str(e)[:480]
            db.session.commit()
        except Exception:
            db.session.rollback()
        return False, str(e)


def build_rag_context_v2(question: str, top_k: int = 5) -> Tuple[str, List[Dict]]:
    """新版 RAG：返回 (上下文文本, citations 列表)
    citations 用于前后端引文展示。
    """
    hits = search_chunks(question, top_k=top_k, max_per_doc=2)
    if not hits:
        return "", []

    lines = ["【知识库检索】（请基于以下资料回答；引用时用 [1][2] 等角标对应到下方资料编号）"]
    citations = []
    for i, h in enumerate(hits, start=1):
        page_info = f"·第{h['page_no']}页" if h.get("page_no") else ""
        lines.append(f"[{i}] {h['doc_title']}{page_info}")
        # 取片段（避免过长撑爆 prompt）
        snippet = (h["content"] or "").strip()
        if len(snippet) > 480:
            snippet = snippet[:480] + "…"
        lines.append(snippet)
        lines.append("")
        citations.append({
            "ref": i,
            "doc_id": h["doc_id"],
            "chunk_id": h["chunk_id"],
            "doc_title": h["doc_title"],
            "doc_type": h["doc_type"],
            "page_no": h.get("page_no"),
            "score": h["score"],
            "snippet": snippet[:120],   # 前台只看 120 字预览
        })
    return "\n".join(lines), citations
